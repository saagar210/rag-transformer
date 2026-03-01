import io
import json
import pytest
from unittest.mock import AsyncMock, MagicMock
from starlette.testclient import TestClient
import app as app_module
from app import app


# ── TestClient fixture ──────────────────────────────────────────
@pytest.fixture
def client():
    with TestClient(app) as c:
        yield c


# ── Mock Ollama client (autouse) ────────────────────────────────
# Patches AsyncClient at class level so the lifespan gets our mock
# when it calls ollama_client = AsyncClient().
@pytest.fixture(autouse=True)
def mock_ollama(monkeypatch):
    mock = AsyncMock()

    # Default: list() returns empty model list
    list_resp = MagicMock()
    list_resp.models = []
    mock.list = AsyncMock(return_value=list_resp)

    # Default: chat() returns an async iterator with one token
    async def default_chat_gen():
        chunk = MagicMock()
        chunk.message.content = "Hello"
        yield chunk

    mock.chat = AsyncMock(return_value=default_chat_gen())

    # Patch the class so the lifespan instantiation gives us our mock
    monkeypatch.setattr("app.AsyncClient", lambda: mock)
    # Also patch the module variable directly in case it was already set
    monkeypatch.setattr(app_module, "ollama_client", mock)

    return mock


# ── DOCX bytes fixture ──────────────────────────────────────────
@pytest.fixture
def sample_docx_bytes():
    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── PDF bytes fixture ───────────────────────────────────────────
@pytest.fixture
def sample_pdf_bytes():
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello from PDF")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── SSE parse helper ────────────────────────────────────────────
def parse_sse(text: str) -> list[dict]:
    events = []
    for line in text.splitlines():
        if line.startswith("data: "):
            events.append(json.loads(line[6:]))
    return events
